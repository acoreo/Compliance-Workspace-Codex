"""Text extraction from evidence files (P3-T01).

Supports PDF, DOCX, XLSX/CSV, plain text, EML, MSG, image (OCR), and a UTF-8
fallback for all other formats.  Results are cached in the evidence_text
SQLite table so that re-extraction is skipped on subsequent runs.

EML and MSG files also yield AttachmentResult objects for each attachment found
inside the message; these are persisted to the email_attachments table.
"""
from __future__ import annotations

import io
import os
import sqlite3
import tempfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional


@dataclass
class AttachmentResult:
    parent_file_path: str       # path to the .eml or .msg that contained this
    attachment_filename: str    # original filename from the email
    attachment_ext: str         # lowercase extension, no dot
    content_bytes: bytes        # raw bytes of the attachment
    text: str                   # extracted text (empty if extraction failed)
    extraction_method: str
    error: Optional[str]
    size_bytes: int


@dataclass
class ExtractResult:
    file_node_id: int
    file_path: str
    text: str
    char_count: int
    extraction_method: str   # "pdfminer" | "python-docx" | "openpyxl" | "plain" | "fallback"
                             # | "email_stdlib" | "extract_msg" | "msg_fallback"
                             # | "pytesseract" | "ocr_unavailable"
    error: Optional[str]
    attachments: list = field(default_factory=list)  # list[AttachmentResult]


# ---------------------------------------------------------------------------
# Attachment sub-extractor
# ---------------------------------------------------------------------------

def _extract_attachment(parent_path: str, filename: str, content_bytes: bytes) -> AttachmentResult:
    """Extract text from raw attachment bytes, routing by file extension."""
    ext = Path(filename).suffix.lower().lstrip(".")
    text = ""
    method = "unsupported"
    error: Optional[str] = None

    if ext == "pdf":
        try:
            from pdfminer.high_level import extract_text as _pdf_extract  # type: ignore
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tf:
                tf.write(content_bytes)
                tf_path = tf.name
            try:
                text = _pdf_extract(tf_path) or ""
                method = "pdfminer"
            finally:
                os.unlink(tf_path)
        except ImportError:
            method = "fallback"
            error = "pdfminer not installed"
        except Exception as exc:
            method = "fallback"
            error = str(exc)

    elif ext in ("docx", "doc"):
        try:
            import docx  # type: ignore
            doc = docx.Document(io.BytesIO(content_bytes))
            text = "\n".join(p.text for p in doc.paragraphs)
            method = "python-docx"
        except ImportError:
            method = "fallback"
            error = "python-docx not installed"
        except Exception as exc:
            method = "fallback"
            error = str(exc)

    elif ext in ("xlsx", "xls"):
        try:
            import openpyxl  # type: ignore
            wb = openpyxl.load_workbook(io.BytesIO(content_bytes), data_only=True)
            rows: list[str] = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    rows.append("\t".join("" if c is None else str(c) for c in row))
            text = "\n".join(rows)
            method = "openpyxl"
        except ImportError:
            method = "fallback"
            error = "openpyxl not installed"
        except Exception as exc:
            method = "fallback"
            error = str(exc)

    elif ext in ("txt", "csv"):
        try:
            text = content_bytes.decode("utf-8", errors="replace")
            method = "plain"
        except Exception as exc:
            method = "fallback"
            error = str(exc)

    elif ext in ("png", "jpg", "jpeg"):
        try:
            from PIL import Image  # type: ignore
            import pytesseract  # type: ignore
            text = pytesseract.image_to_string(Image.open(io.BytesIO(content_bytes)))
            method = "pytesseract"
        except ImportError:
            method = "ocr_unavailable"
            error = (
                "pytesseract not installed — on Windows install via: "
                "winget install UB-Mannheim.TesseractOCR  "
                "(or download from https://github.com/UB-Mannheim/tesseract/wiki), "
                "then: pip install Pillow pytesseract"
            )
        except Exception as exc:
            method = "ocr_unavailable"
            error = str(exc)

    return AttachmentResult(
        parent_file_path=parent_path,
        attachment_filename=filename,
        attachment_ext=ext,
        content_bytes=content_bytes,
        text=text,
        extraction_method=method,
        error=error,
        size_bytes=len(content_bytes),
    )


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------

def _extract_pdf(path: Path) -> tuple[str, str, Optional[str]]:
    try:
        from pdfminer.high_level import extract_text  # type: ignore
        text = extract_text(str(path)) or ""
        return text, "pdfminer", None
    except Exception as exc:
        try:
            raw = path.read_bytes().decode("utf-8", errors="replace")
            return raw, "fallback", str(exc)
        except Exception as exc2:
            return "", "fallback", str(exc2)


def _extract_docx(path: Path) -> tuple[str, str, Optional[str]]:
    # Primary: python-docx
    try:
        import docx  # type: ignore
        doc = docx.Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        return text, "python-docx", None
    except ImportError:
        pass
    except Exception:
        pass

    # Fallback: zipfile + strip XML tags
    try:
        import re
        import zipfile

        parts: list[str] = []
        with zipfile.ZipFile(str(path)) as zf:
            for name in zf.namelist():
                if name.startswith("word/") and name.endswith(".xml"):
                    xml = zf.read(name).decode("utf-8", errors="replace")
                    parts.append(re.sub(r"<[^>]+>", " ", xml))
        return " ".join(parts), "fallback", None
    except Exception as exc:
        return path.read_bytes().decode("utf-8", errors="replace"), "fallback", str(exc)


def _extract_xlsx(path: Path) -> tuple[str, str, Optional[str]]:
    # Primary: openpyxl
    try:
        import openpyxl  # type: ignore

        wb = openpyxl.load_workbook(str(path), data_only=True)
        rows: list[str] = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                rows.append("\t".join("" if c is None else str(c) for c in row))
        return "\n".join(rows), "openpyxl", None
    except ImportError:
        pass
    except Exception:
        pass

    # Fallback: plain read
    try:
        return path.read_text(encoding="utf-8", errors="replace"), "plain", None
    except Exception as exc:
        return "", "fallback", str(exc)


def _extract_csv(path: Path) -> tuple[str, str, Optional[str]]:
    try:
        import csv as _csv

        rows: list[str] = []
        with path.open(newline="", encoding="utf-8", errors="replace") as f:
            reader = _csv.reader(f)
            for row in reader:
                rows.append("\t".join(row))
        return "\n".join(rows), "plain", None
    except Exception as exc:
        return "", "fallback", str(exc)


def _extract_eml(path: Path) -> tuple[str, str, Optional[str], list]:
    """Extract body text and attachments from an .eml file.

    Returns (body_text, method, error, attachments).
    """
    import email as _email
    import re

    try:
        raw = path.read_bytes()
        msg = _email.message_from_bytes(raw)

        headers: list[str] = []
        for header in ("Subject", "From", "To", "Date"):
            val = msg.get(header, "")
            if val:
                headers.append(f"{header}: {val}")

        body_parts: list[str] = []
        html_parts: list[str] = []
        attachments: list[AttachmentResult] = []

        for part in msg.walk():
            ct = part.get_content_type()
            filename = part.get_filename()
            disp = str(part.get("Content-Disposition", "") or "")

            # Attachment: explicit filename or Content-Disposition: attachment
            if filename or "attachment" in disp.lower():
                payload = part.get_payload(decode=True)
                if payload:
                    att = _extract_attachment(
                        str(path), filename or "attachment", payload
                    )
                    attachments.append(att)
            elif ct == "text/plain":
                charset = part.get_content_charset() or "utf-8"
                payload = part.get_payload(decode=True)
                if payload:
                    body_parts.append(payload.decode(charset, errors="replace"))
            elif ct == "text/html":
                charset = part.get_content_charset() or "utf-8"
                payload = part.get_payload(decode=True)
                if payload:
                    html = payload.decode(charset, errors="replace")
                    html_parts.append(re.sub(r"<[^>]+>", " ", html))

        body = "\n".join(body_parts) if body_parts else "\n".join(html_parts)
        text = "\n".join(headers) + ("\n\n" + body if body else "")
        return text.strip(), "email_stdlib", None, attachments
    except Exception as exc:
        return "", "fallback", str(exc), []


def _extract_msg(path: Path) -> tuple[str, str, Optional[str], list]:
    """Extract body text and attachments from an Outlook .msg file.

    Returns (body_text, method, error, attachments).
    Requires the extract-msg library for attachment extraction; without it,
    falls back to a raw ASCII scan (no attachment data available).
    """
    # Primary: extract-msg library
    try:
        import extract_msg  # type: ignore

        with extract_msg.Message(str(path)) as m:
            body = m.body or ""
            attachments: list[AttachmentResult] = []
            for att in (m.attachments or []):
                filename = att.longFilename or att.shortFilename or "attachment"
                data = att.data
                if data:
                    attachments.append(_extract_attachment(str(path), filename, data))
        return body, "extract_msg", None, attachments
    except ImportError:
        pass
    except Exception:
        pass

    # Fallback: scan raw bytes for readable strings
    import re

    try:
        raw = path.read_bytes()
        # Extract ASCII/UTF-8 runs longer than 20 characters
        candidates = re.findall(rb"[\x20-\x7E]{20,}", raw)
        text = "\n".join(c.decode("ascii", errors="replace") for c in candidates)
        return (
            text,
            "msg_fallback",
            "full attachment extraction requires the extract-msg library",
            [],
        )
    except Exception as exc:
        return "", "fallback", str(exc), []


def _extract_image(path: Path) -> tuple[str, str, Optional[str]]:
    try:
        from PIL import Image  # type: ignore
        import pytesseract  # type: ignore

        text = pytesseract.image_to_string(Image.open(str(path)))
        return text, "pytesseract", None
    except ImportError:
        return (
            "",
            "ocr_unavailable",
            (
                "pytesseract not installed — on Windows install via: "
                "winget install UB-Mannheim.TesseractOCR  "
                "(or download from https://github.com/UB-Mannheim/tesseract/wiki), "
                "then: pip install Pillow pytesseract"
            ),
        )
    except Exception as exc:
        return "", "ocr_unavailable", str(exc)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_file(file_node_id: int, file_path: str) -> ExtractResult:
    """Extract text from a single file, choosing method by extension."""
    path = Path(file_path)
    if not path.exists():
        return ExtractResult(
            file_node_id=file_node_id,
            file_path=file_path,
            text="",
            char_count=0,
            extraction_method="fallback",
            error=f"File not found: {file_path}",
        )

    ext = path.suffix.lower().lstrip(".")
    attachments: list = []

    if ext == "pdf":
        text, method, error = _extract_pdf(path)
    elif ext in ("docx", "doc"):
        text, method, error = _extract_docx(path)
    elif ext in ("xlsx", "xls"):
        text, method, error = _extract_xlsx(path)
    elif ext == "csv":
        text, method, error = _extract_csv(path)
    elif ext == "eml":
        text, method, error, attachments = _extract_eml(path)
    elif ext == "msg":
        text, method, error, attachments = _extract_msg(path)
    elif ext in ("png", "jpg", "jpeg"):
        text, method, error = _extract_image(path)
    elif ext == "txt":
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
            method, error = "plain", None
        except Exception as exc:
            text, method, error = "", "fallback", str(exc)
    else:
        try:
            text = path.read_bytes().decode("utf-8", errors="replace")
            method, error = "fallback", None
        except Exception as exc:
            text, method, error = "", "fallback", str(exc)

    return ExtractResult(
        file_node_id=file_node_id,
        file_path=file_path,
        text=text,
        char_count=len(text),
        extraction_method=method,
        error=error,
        attachments=attachments,
    )


def extract_all(
    conn: sqlite3.Connection,
    scan_id: int,
    *,
    force: bool = False,
) -> list[ExtractResult]:
    """Extract text from all document file_nodes for scan_id.

    Results are cached in evidence_text; already-cached files are skipped
    unless *force* is True.  Attachment metadata is written to email_attachments.
    """
    rows = conn.execute(
        "SELECT id, full_path FROM file_nodes WHERE scan_id = ? AND is_document = 1",
        (scan_id,),
    ).fetchall()

    results: list[ExtractResult] = []
    for file_node_id, file_path in rows:
        if not force:
            cached = conn.execute(
                "SELECT text, extraction_method, char_count, error "
                "FROM evidence_text WHERE file_node_id = ?",
                (file_node_id,),
            ).fetchone()
            if cached:
                results.append(
                    ExtractResult(
                        file_node_id=file_node_id,
                        file_path=file_path,
                        text=cached[0],
                        char_count=cached[2],
                        extraction_method=cached[1],
                        error=cached[3],
                    )
                )
                continue

        result = extract_file(file_node_id, file_path)
        _cache_result(conn, result)
        if result.attachments:
            _cache_attachments(conn, file_node_id, result.attachments)
        results.append(result)

    return results


def _cache_result(conn: sqlite3.Connection, result: ExtractResult) -> None:
    ts = datetime.now(timezone.utc).isoformat()
    conn.execute(
        """INSERT OR REPLACE INTO evidence_text
               (file_node_id, text, extraction_method, char_count, extracted_ts, error)
               VALUES (?, ?, ?, ?, ?, ?)""",
        (
            result.file_node_id,
            result.text,
            result.extraction_method,
            result.char_count,
            ts,
            result.error,
        ),
    )
    conn.commit()


def _cache_attachments(
    conn: sqlite3.Connection,
    file_node_id: int,
    attachments: list,
) -> None:
    """Persist attachment extraction results to email_attachments table."""
    ts = datetime.now(timezone.utc).isoformat()
    for att in attachments:
        conn.execute(
            """INSERT INTO email_attachments
                   (file_node_id, filename, extension, size_bytes, text,
                    extraction_method, error, extracted_ts)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                file_node_id,
                att.attachment_filename,
                att.attachment_ext,
                att.size_bytes,
                att.text,
                att.extraction_method,
                att.error,
                ts,
            ),
        )
    conn.commit()
